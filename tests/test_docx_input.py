"""DOCX *input* end-to-end test.

Unlike DOCX output (which we build ourselves), DOCX input must be rendered to
PDF by LibreOffice to obtain real page coordinates. So this test is skipped
automatically when LibreOffice isn't installed, rather than failing on machines
without it.

It builds a .docx with python-docx (heading + paragraph + a bordered table),
runs the full pipeline, and asserts the handwriting output preserves the page,
the table borders (vector graphics), and renders text in the handwriting font.
"""

from __future__ import annotations

import os
import sys

import pytest

_SRC = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "src"))
if _SRC not in sys.path:
    sys.path.insert(0, _SRC)

import fitz  # noqa: E402

from handscrybe.config import Config  # noqa: E402
from handscrybe.normalize import find_soffice  # noqa: E402
from handscrybe.pipeline import convert  # noqa: E402

# Skip the whole module if LibreOffice isn't available.
pytestmark = pytest.mark.skipif(
    find_soffice() is None,
    reason="LibreOffice (soffice) not installed; DOCX input requires it.",
)

_HEADING = "Project Report"
_BODY = "This document verifies DOCX input works end to end."


def _make_docx(path: str) -> None:
    from docx import Document

    d = Document()
    d.add_heading(_HEADING, level=1)
    d.add_paragraph(_BODY)
    table = d.add_table(rows=2, cols=2)
    table.style = "Table Grid"  # gives the table visible borders
    for r in range(2):
        for c in range(2):
            table.cell(r, c).text = f"cell {r},{c}"
    d.save(path)


def _count_drawings(path: str) -> int:
    doc = fitz.open(path)
    try:
        return sum(len(doc[p].get_drawings()) for p in range(doc.page_count))
    finally:
        doc.close()


def test_docx_input_end_to_end(tmp_path):
    src = str(tmp_path / "sample.docx")
    out = str(tmp_path / "sample_hand.pdf")
    _make_docx(src)

    convert(src, out, Config(font_bold=None))

    assert os.path.isfile(out)
    doc = fitz.open(out)
    try:
        assert doc.page_count >= 1
        # Table borders (vector graphics) survived.
        assert _count_drawings(out) > 0
        # Text renders in the handwriting font, and the content is preserved.
        text = "\n".join(doc[p].get_text() for p in range(doc.page_count))
        assert _HEADING in text
        fonts = {
            span.get("font", "")
            for p in range(doc.page_count)
            for block in doc[p].get_text("dict")["blocks"]
            for line in block.get("lines", [])
            for span in line["spans"]
            if span.get("text", "").strip()
        }
        assert any("caveat" in f.lower() for f in fonts)
    finally:
        doc.close()
