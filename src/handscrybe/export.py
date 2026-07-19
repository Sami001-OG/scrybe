"""Deliver the conversion result in the format the user asked for.

The pipeline always produces a handwriting PDF first (that's the native,
highest-fidelity artifact). This module turns that PDF — plus the normalized
source PDF, which still carries an extractable text layer — into whatever the
user wants:

    PDF  -> the handwriting PDF itself (a copy to the requested path).
    DOCX -> the handwriting PDF converted back through LibreOffice, so every
            page becomes a full-page handwriting image inside a .docx.
    TXT  -> the document's plain text (handwriting can't live in a text file,
            so we deliver the CONTENT instead).
    MD   -> the same text with light Markdown structure (paragraphs split on
            blank lines).

WHY TEXT COMES FROM THE SOURCE PDF, NOT THE HANDWRITING PDF
----------------------------------------------------------
The handwriting PDF's text is either drawn as images (user-glyph mode) or as a
handwriting font whose extracted characters are still the real letters — but
the *source* (normalized) PDF is the clean, authoritative text layer. So for
TXT/MD we extract from the source PDF. That keeps the text output correct
regardless of how the handwriting was rendered.
"""

from __future__ import annotations

import os
import shutil

import fitz  # PyMuPDF

from .config import OutputFormat


def _extract_text_by_page(pdf_path: str) -> list[str]:
    """Return the plain text of each page of a PDF, in order."""
    doc = fitz.open(pdf_path)
    try:
        return [doc[p].get_text("text") for p in range(doc.page_count)]
    finally:
        doc.close()


def _pdf_to_docx(handwriting_pdf: str, out_path: str, soffice_cmd: str | None) -> str:
    """Build a DOCX that shows each handwriting PDF page as a full-page image.

    We deliberately do NOT route this through LibreOffice: LibreOffice opens a
    PDF in its Draw module, which has no Writer/DOCX export filter, so a direct
    PDF->DOCX conversion aborts with "no export filter". Instead we rasterize
    each PDF page with PyMuPDF and drop it, sized to the page's text width, into
    a python-docx document whose page size and margins mirror the source. The
    result is a real, openable .docx that preserves the handwriting visually and
    keeps the original pagination (one source page per Word page).

    This has no external-tool dependency, which also makes DOCX output work on
    machines without LibreOffice.
    """
    from docx import Document as DocxDocument
    from docx.shared import Emu

    # Rasterize at ~150 DPI: crisp enough for handwriting, small enough to keep
    # the .docx reasonable. 150/72 is the zoom over PDF's native 72 DPI.
    zoom = 150.0 / 72.0
    doc = DocxDocument()
    src = fitz.open(handwriting_pdf)
    try:
        # EMUs are the DOCX unit; 914400 EMU per inch, PDF points are 1/72 inch.
        emu_per_point = 914400.0 / 72.0
        first = True
        for pno in range(src.page_count):
            page = src[pno]
            rect = page.rect
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            img_bytes = pix.tobytes("png")

            section = doc.sections[0] if first else doc.add_section()
            # Match the Word page to the PDF page so pagination is 1:1 and the
            # image fills the sheet with zero margins (the PDF already contains
            # its own whitespace/margins).
            section.page_width = Emu(int(rect.width * emu_per_point))
            section.page_height = Emu(int(rect.height * emu_per_point))
            for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
                setattr(section, m, Emu(0))

            para = doc.add_paragraph()
            para.paragraph_format.space_after = Emu(0)
            run = para.add_run()
            import io

            run.add_picture(io.BytesIO(img_bytes), width=Emu(int(rect.width * emu_per_point)))
            first = False
    finally:
        src.close()

    doc.save(out_path)
    return out_path


def _text_to_markdown(pages: list[str]) -> str:
    """Turn extracted per-page text into light Markdown.

    We keep it deliberately simple and lossless-leaning: pages are separated by
    a horizontal rule, and runs of text separated by blank lines become
    paragraphs. We don't guess headings/lists from font sizes here — that would
    risk mangling content; the goal is faithful text with basic structure."""
    parts: list[str] = []
    for i, text in enumerate(pages):
        if i > 0:
            parts.append("\n\n---\n\n")  # page break as a horizontal rule
        # Normalize line endings and collapse 3+ blank lines to a paragraph gap.
        lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        buf: list[str] = []
        blank = 0
        for ln in lines:
            if ln.strip() == "":
                blank += 1
                if blank == 1:
                    buf.append("")  # single paragraph separator
            else:
                blank = 0
                buf.append(ln.rstrip())
        parts.append("\n".join(buf).strip())
    return "\n".join(parts).strip() + "\n"


def deliver(
    handwriting_pdf: str,
    source_pdf: str,
    output_path: str,
    fmt: OutputFormat,
    soffice_cmd: str | None = None,
) -> str:
    """Produce `output_path` in format `fmt` from the conversion artifacts.

    - `handwriting_pdf`: the rendered handwriting PDF (visual result).
    - `source_pdf`: the normalized source PDF (clean text layer, for TXT/MD).
    Returns `output_path`.
    """
    if fmt is OutputFormat.PDF:
        if os.path.abspath(handwriting_pdf) != os.path.abspath(output_path):
            shutil.copyfile(handwriting_pdf, output_path)
        return output_path

    if fmt is OutputFormat.DOCX:
        return _pdf_to_docx(handwriting_pdf, output_path, soffice_cmd)

    # Text families: extract from the authoritative source text layer.
    pages = _extract_text_by_page(source_pdf)
    if fmt is OutputFormat.TXT:
        # Join pages with a form-feed so page boundaries are recoverable but the
        # file stays plain text.
        content = "\f".join(p.rstrip() for p in pages).strip() + "\n"
        with open(output_path, "w", encoding="utf-8") as fh:
            fh.write(content)
        return output_path

    if fmt is OutputFormat.MD:
        with open(output_path, "w", encoding="utf-8") as fh:
            fh.write(_text_to_markdown(pages))
        return output_path

    raise ValueError(f"Unsupported output format: {fmt!r}")
